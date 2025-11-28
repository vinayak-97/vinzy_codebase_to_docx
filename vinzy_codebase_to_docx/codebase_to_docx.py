"""
codebase_to_docx: Convert your entire codebase to a Word document
"""

import os
import sys
from pathlib import Path
from typing import List, Optional, Set

# Fix for docx import issues
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError as e:
    print("ERROR: python-docx is not installed correctly.")
    print("\nPlease run these commands:")
    print("  pip uninstall docx python-docx -y")
    print("  pip install python-docx")
    print(f"\nOriginal error: {e}")
    sys.exit(1)

import mimetypes

class CodebaseConverter:
    """Convert a codebase directory to a Word document."""
    
    # Common code file extensions
    CODE_EXTENSIONS = {
        '.py', '.js', '.jsx', '.ts', '.tsx', '.java', '.cpp', '.c', '.h',
        '.cs', '.rb', '.go', '.rs', '.php', '.swift', '.kt', '.scala',
        '.html', '.css', '.scss', '.sass', '.less', '.xml', '.json',
        '.yaml', '.yml', '.md', '.txt', '.sh', '.bash', '.sql', '.r'
    }
    
    # Directories to ignore by default
    IGNORE_DIRS = {
        '__pycache__', 'node_modules', '.git', '.venv', 'venv',
        'env', 'build', 'dist', '.idea', '.vscode', 'target',
        '.gradle', 'bin', 'obj', 'coverage'
    }
    
    def __init__(self, 
                 codebase_path: str,
                 output_path: str = 'codebase.docx',
                 include_file_paths: bool = True,
                 include_toc: bool = True,
                 ignore_dirs: Optional[Set[str]] = None,
                 include_extensions: Optional[Set[str]] = None):
        """
        Initialize the converter.
        
        Args:
            codebase_path: Path to the codebase directory
            output_path: Path for the output Word document
            include_file_paths: Whether to include file paths in the document
            include_toc: Whether to include a table of contents
            ignore_dirs: Additional directories to ignore
            include_extensions: Specific file extensions to include (None = use defaults)
        """
        self.codebase_path = Path(codebase_path).resolve()
        self.output_path = output_path
        self.include_file_paths = include_file_paths
        self.include_toc = include_toc
        
        # Setup ignore directories
        self.ignore_dirs = self.IGNORE_DIRS.copy()
        if ignore_dirs:
            self.ignore_dirs.update(ignore_dirs)
        
        # Setup file extensions
        self.include_extensions = include_extensions or self.CODE_EXTENSIONS
        
        # Validate paths
        if not self.codebase_path.exists():
            raise ValueError(f"Codebase path does not exist: {codebase_path}")
        if not self.codebase_path.is_dir():
            raise ValueError(f"Codebase path is not a directory: {codebase_path}")
    
    def _should_include_file(self, file_path: Path) -> bool:
        """Check if a file should be included in the document."""
        # Check extension
        if file_path.suffix.lower() not in self.include_extensions:
            return False
        
        # Check if in ignored directory
        for parent in file_path.parents:
            if parent.name in self.ignore_dirs:
                return False
        
        return True
    
    def _get_files(self) -> List[Path]:
        """Get all files to include in the document."""
        files = []
        for root, dirs, filenames in os.walk(self.codebase_path):
            # Remove ignored directories from dirs list
            dirs[:] = [d for d in dirs if d not in self.ignore_dirs]
            
            for filename in filenames:
                file_path = Path(root) / filename
                if self._should_include_file(file_path):
                    files.append(file_path)
        
        return sorted(files)
    
    def _add_title(self, doc: Document, title: str):
        """Add a title to the document."""
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def _add_table_of_contents(self, doc: Document, files: List[Path]):
        """Add a table of contents listing all files."""
        doc.add_heading('Table of Contents', level=1)
        
        for file_path in files:
            rel_path = file_path.relative_to(self.codebase_path)
            p = doc.add_paragraph(str(rel_path), style='List Bullet')
            # Make it smaller
            for run in p.runs:
                run.font.size = Pt(10)
        
        doc.add_page_break()
    
    def _add_file_content(self, doc: Document, file_path: Path):
        """Add a single file's content to the document."""
        rel_path = file_path.relative_to(self.codebase_path)
        
        # Add file heading
        doc.add_heading(str(rel_path), level=1)
        
        # Add file path if requested
        if self.include_file_paths:
            p = doc.add_paragraph()
            run = p.add_run(f"Location: {file_path}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.italic = True
        
        # Read and add file content
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Add code block
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            
            # Add light gray background (simulating code block)
            # Note: python-docx doesn't support full background shading,
            # but we can make the font distinctive
            
        except Exception as e:
            p = doc.add_paragraph()
            run = p.add_run(f"Error reading file: {str(e)}")
            run.font.color.rgb = RGBColor(255, 0, 0)
        
        doc.add_paragraph()  # Add spacing
    
    def convert(self, progress_callback=None) -> str:
        """
        Convert the codebase to a Word document.
        
        Args:
            progress_callback: Optional callback function(current, total, filename)
        
        Returns:
            Path to the created document
        """
        print(f"Scanning codebase: {self.codebase_path}")
        files = self._get_files()
        
        if not files:
            raise ValueError("No files found to convert!")
        
        print(f"Found {len(files)} files to convert")
        
        # Create document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Add title
        self._add_title(doc, f"Codebase: {self.codebase_path.name}")
        doc.add_paragraph(f"Total files: {len(files)}")
        doc.add_page_break()
        
        # Add table of contents
        if self.include_toc:
            self._add_table_of_contents(doc, files)
        
        # Add each file
        for idx, file_path in enumerate(files, 1):
            if progress_callback:
                progress_callback(idx, len(files), file_path.name)
            
            print(f"Processing [{idx}/{len(files)}]: {file_path.name}")
            self._add_file_content(doc, file_path)
            
            # Add page break between files (except last one)
            if idx < len(files):
                doc.add_page_break()
        
        # Save document
        doc.save(self.output_path)
        print(f"\nDocument saved: {self.output_path}")
        
        return self.output_path
    
    def author(self) -> str:
        """Return the author of the codebase converter."""
        return "Developed by Vinayak Bhosale"


# Example usage function
def convert_codebase(codebase_path: str,
                    output_path: str = 'codebase.docx',
                    include_file_paths: bool = True,
                    include_toc: bool = True,
                    ignore_dirs: Optional[Set[str]] = None,
                    include_extensions: Optional[Set[str]] = None):
    """
    Convert a codebase to a Word document.
    
    Args:
        codebase_path: Path to the codebase directory
        output_path: Path for the output Word document
        include_file_paths: Whether to include file paths in the document
        include_toc: Whether to include table of contents
        ignore_dirs: Additional directories to ignore
        include_extensions: Specific file extensions to include
    
    Returns:
        Path to the created document
    """
    converter = CodebaseConverter(
        codebase_path=codebase_path,
        output_path=output_path,
        include_file_paths=include_file_paths,
        include_toc=include_toc,
        ignore_dirs=ignore_dirs,
        include_extensions=include_extensions
    )
    
    return converter.convert()


if __name__ == "__main__":
    # Example usage
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python codebase_to_docx.py <codebase_path> [output_path]")
        sys.exit(1)
    
    codebase_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else 'codebase.docx'
    
    try:
        convert_codebase(
            codebase_path=codebase_path,
            output_path=output_path,
            include_file_paths=True,
            include_toc=True
        )
        print("\nConversion completed successfully!")
    except Exception as e:
        print(f"\nError: {str(e)}")
        sys.exit(1)